"""
Multi-format Document Extractor

Supports extraction from:
- PDF (existing)
- Word (.docx, .doc)
- Excel (.xlsx, .xls)
"""

import logging
from pathlib import Path
from typing import Dict, List, Optional
from dataclasses import dataclass

# Import existing PDF extractor
from .pdf_to_text import extract_text_from_pdf

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


@dataclass
class DocumentContent:
    """Container for extracted document content."""
    text: str
    pages: List[Dict]
    total_pages: int
    total_chars: int
    extraction_method: str
    file_type: str


class MultiFormatExtractor:
    """
    Extract text from multiple document formats.
    
    Supported formats:
    - PDF (.pdf)
    - Word (.docx, .doc)
    - Excel (.xlsx, .xls)
    """
    
    SUPPORTED_FORMATS = {
        'pdf': ['.pdf'],
        'word': ['.docx', '.doc'],
        'excel': ['.xlsx', '.xls']
    }
    
    def __init__(self):
        """Initialize extractor."""
        self.extractors = {
            'pdf': self._extract_pdf,
            'word': self._extract_word,
            'excel': self._extract_excel
        }
    
    def extract_document(self, file_path: str) -> DocumentContent:
        """
        Extract text from any supported document format.
        
        Args:
            file_path: Path to document
            
        Returns:
            DocumentContent with extracted text
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Detect file type
        file_type = self._detect_file_type(file_path)
        
        if file_type not in self.extractors:
            raise ValueError(
                f"Unsupported file format: {file_path.suffix}. "
                f"Supported: {self._get_supported_extensions()}"
            )
        
        logger.info(f"Extracting {file_type.upper()} file: {file_path.name}")
        
        # Extract using appropriate method
        return self.extractors[file_type](file_path)
    
    def _detect_file_type(self, file_path: Path) -> str:
        """Detect document type from extension."""
        ext = file_path.suffix.lower()
        
        for file_type, extensions in self.SUPPORTED_FORMATS.items():
            if ext in extensions:
                return file_type
        
        return None
    
    def _get_supported_extensions(self) -> List[str]:
        """Get list of all supported extensions."""
        extensions = []
        for exts in self.SUPPORTED_FORMATS.values():
            extensions.extend(exts)
        return extensions
    
    def _extract_pdf(self, file_path: Path) -> DocumentContent:
        """Extract text from PDF using existing extractor."""
        result = extract_text_from_pdf(str(file_path))
        
        return DocumentContent(
            text=result['full_text'],
            pages=result['pages'],
            total_pages=result['total_pages'],
            total_chars=result['total_chars'],
            extraction_method=result['extraction_method'],
            file_type='pdf'
        )
    
    def _extract_word(self, file_path: Path) -> DocumentContent:
        """
        Extract text from Word documents.
        
        Uses python-docx for .docx files.
        """
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx required. Install: pip install python-docx")
        
        try:
            # Load document
            doc = Document(str(file_path))
            
            # Extract all paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Extract tables if any
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text)
            
            full_text = '\n\n'.join(paragraphs)
            
            # Estimate pages (rough approximation)
            # Assume ~500 chars per page
            estimated_pages = max(1, len(full_text) // 500)
            
            # Create page info
            pages = [{
                'page_number': i + 1,
                'text': full_text[i*500:(i+1)*500],
                'is_scanned': False,
                'char_count': min(500, len(full_text) - i*500)
            } for i in range(estimated_pages)]
            
            logger.info(f"Extracted {len(paragraphs)} paragraphs from Word document")
            
            return DocumentContent(
                text=full_text,
                pages=pages,
                total_pages=estimated_pages,
                total_chars=len(full_text),
                extraction_method='python-docx',
                file_type='word'
            )
            
        except Exception as e:
            logger.error(f"Word extraction failed: {e}")
            raise
    
    def _extract_excel(self, file_path: Path) -> DocumentContent:
        """
        Extract text from Excel spreadsheets.
        
        Uses openpyxl for .xlsx files and xlrd for .xls files.
        Extracts all sheets, all cells.
        """
        ext = file_path.suffix.lower()
        
        if ext == '.xlsx':
            return self._extract_xlsx(file_path)
        elif ext == '.xls':
            return self._extract_xls(file_path)
        else:
            raise ValueError(f"Unsupported Excel format: {ext}")
    
    def _extract_xlsx(self, file_path: Path) -> DocumentContent:
        """Extract from .xlsx files using openpyxl."""
        try:
            from openpyxl import load_workbook
        except ImportError:
            raise ImportError("openpyxl required. Install: pip install openpyxl")
        
        try:
            # Load workbook
            wb = load_workbook(str(file_path), data_only=True)
            
            all_text = []
            sheet_count = 0
            
            # Extract from all sheets
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                sheet_count += 1
                
                # Add sheet header
                all_text.append(f"\n=== Sheet: {sheet_name} ===\n")
                
                # Extract all cells
                for row in sheet.iter_rows(values_only=True):
                    # Filter out None values and convert to strings
                    cell_values = [str(cell) for cell in row if cell is not None]
                    if cell_values:
                        all_text.append(' | '.join(cell_values))
            
            full_text = '\n'.join(all_text)
            
            # Create "pages" as sheets
            pages = [{
                'page_number': i + 1,
                'text': f"Sheet {i+1}",
                'is_scanned': False,
                'char_count': len(full_text) // sheet_count
            } for i in range(sheet_count)]
            
            logger.info(f"Extracted {sheet_count} sheets from Excel file")
            
            return DocumentContent(
                text=full_text,
                pages=pages,
                total_pages=sheet_count,
                total_chars=len(full_text),
                extraction_method='openpyxl',
                file_type='excel'
            )
            
        except Exception as e:
            logger.error(f"Excel (.xlsx) extraction failed: {e}")
            raise
    
    def _extract_xls(self, file_path: Path) -> DocumentContent:
        """Extract from legacy .xls files using xlrd."""
        try:
            import xlrd
        except ImportError:
            raise ImportError("xlrd required. Install: pip install xlrd")
        
        try:
            # Load workbook
            wb = xlrd.open_workbook(str(file_path))
            
            all_text = []
            sheet_count = wb.nsheets
            
            # Extract from all sheets
            for sheet_idx in range(sheet_count):
                sheet = wb.sheet_by_index(sheet_idx)
                
                # Add sheet header
                all_text.append(f"\n=== Sheet: {sheet.name} ===\n")
                
                # Extract all cells
                for row_idx in range(sheet.nrows):
                    row_values = []
                    for col_idx in range(sheet.ncols):
                        cell = sheet.cell(row_idx, col_idx)
                        if cell.value:
                            row_values.append(str(cell.value))
                    
                    if row_values:
                        all_text.append(' | '.join(row_values))
            
            full_text = '\n'.join(all_text)
            
            # Create pages
            pages = [{
                'page_number': i + 1,
                'text': f"Sheet {i+1}",
                'is_scanned': False,
                'char_count': len(full_text) // sheet_count
            } for i in range(sheet_count)]
            
            logger.info(f"Extracted {sheet_count} sheets from legacy Excel file")
            
            return DocumentContent(
                text=full_text,
                pages=pages,
                total_pages=sheet_count,
                total_chars=len(full_text),
                extraction_method='xlrd',
                file_type='excel'
            )
            
        except Exception as e:
            logger.error(f"Excel (.xls) extraction failed: {e}")
            raise


def extract_document(file_path: str) -> DocumentContent:
    """
    Convenience function to extract from any supported format.
    
    Args:
        file_path: Path to document (PDF, Word, or Excel)
        
    Returns:
        DocumentContent with extracted text
    """
    extractor = MultiFormatExtractor()
    return extractor.extract_document(file_path)


if __name__ == "__main__":
    import sys
    
    if len(sys.argv) > 1:
        file_path = sys.argv[1]
        
        try:
            result = extract_document(file_path)
            
            print(f"\n=== Extraction Results ===")
            print(f"File type: {result.file_type.upper()}")
            print(f"Total pages/sheets: {result.total_pages}")
            print(f"Total characters: {result.total_chars}")
            print(f"Method: {result.extraction_method}")
            print(f"\n=== First 500 characters ===")
            print(result.text[:500])
            
        except Exception as e:
            print(f"Error: {e}")
    else:
        print("Usage: python document_extractor.py <file_path>")
        print("\nSupported formats: PDF, Word (.docx, .doc), Excel (.xlsx, .xls)")
